#!/usr/bin/env python3
"""Export shared/static data that rarely changes.

Exports:
- constitution.json - League rules
- hall_of_fame.json - Historical records  
- banners.json - Championship banner images
- transactions.json - All historical transactions
"""

import json
import os
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from . import DOCS_DIR, PROJECT_DIR, SHARED_DIR, WEB_DIR, ensure_dirs


def get_docx_module():
    """Try to import docx module."""
    try:
        import docx
        return docx
    except ImportError:
        return None


def parse_constitution(doc_path: str) -> list[dict]:
    """Parse constitution document into structured sections with nested lists."""
    docx = get_docx_module()
    if not docx:
        return []
    
    doc = docx.Document(doc_path)
    sections = []
    current_article = None
    current_section = None
    
    LEVEL_1 = 800000
    LEVEL_2 = 1200000
    LEVEL_3 = 1600000
    
    def get_indent_level(para):
        left_indent = para.paragraph_format.left_indent
        if left_indent is None:
            return 0
        if left_indent >= LEVEL_3:
            return 3
        if left_indent >= LEVEL_2:
            return 2
        if left_indent >= LEVEL_1:
            return 1
        return 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style = para.style.name if para.style else ''
        
        if style == 'Title':
            continue
        elif style == 'Heading 1':
            current_article = {'title': text, 'sections': []}
            sections.append(current_article)
            current_section = None
        elif style == 'Heading 2':
            if current_article:
                current_section = {'title': text, 'content': []}
                current_article['sections'].append(current_section)
        elif style == 'Heading 3':
            if current_section:
                current_section['content'].append({'type': 'subheader', 'text': text})
        else:
            if current_section:
                indent = get_indent_level(para)
                if indent >= 3:
                    current_section['content'].append({'type': 'subitem', 'text': text})
                elif indent >= 2:
                    current_section['content'].append({'type': 'item', 'text': text})
                else:
                    current_section['content'].append({'type': 'header', 'text': text})
            elif current_article:
                if not current_article.get('intro'):
                    current_article['intro'] = []
                current_article['intro'].append(text)
    
    return sections


def parse_hall_of_fame(doc_path: str) -> dict:
    """Parse Hall of Fame document."""
    docx = get_docx_module()
    if not docx:
        return {}
    
    doc = docx.Document(doc_path)
    
    result = {
        'finishes_by_year': [],
        'mvps': [],
        'team_records': [],
        'player_records': [],
        'owner_stats': [],
    }
    
    current_year = None
    current_section = None
    current_subsection = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style = para.style.name if para.style else ''
        
        if style == 'Title':
            continue
        elif style == 'Heading 1':
            if 'Summary' in text:
                current_section = 'summary'
            elif 'Finishes' in text:
                current_section = 'finishes'
            elif 'Team Records' in text:
                current_section = 'team_records'
            elif 'Player Records' in text:
                current_section = 'player_records'
            else:
                current_section = text
        elif style == 'Heading 2':
            if current_section == 'finishes':
                current_year = {'year': text, 'results': []}
                result['finishes_by_year'].append(current_year)
            elif current_section in ['team_records', 'player_records']:
                if ' over ' in text and '(' in text and ')' in text:
                    if current_subsection:
                        current_subsection['records'].append(text)
                else:
                    current_subsection = {'title': text, 'records': []}
                    result[current_section].append(current_subsection)
            elif 'MVP' in text:
                current_section = 'mvps'
        elif style == 'Heading 3':
            if current_section in ['team_records', 'player_records']:
                current_subsection = {'title': text, 'records': []}
                result[current_section].append(current_subsection)
        else:
            if current_section == 'finishes' and current_year:
                current_year['results'].append(text)
            elif current_section == 'mvps':
                result['mvps'].append(text)
            elif current_subsection:
                current_subsection['records'].append(text)
    
    # Parse owner stats table
    if doc.tables:
        table = doc.tables[0]
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if cells[0]:
                owner_data = dict(zip(headers, cells))
                result['owner_stats'].append(owner_data)
    
    # Clean up empty sections
    result['team_records'] = [s for s in result['team_records'] if s['records']]
    result['player_records'] = [s for s in result['player_records'] if s['records']]
    
    return result


def get_banners() -> list[str]:
    """Get list of banner images."""
    banners_dir = WEB_DIR / "images" / "banners"
    if banners_dir.exists():
        return sorted([f.name for f in banners_dir.glob("*_banner.png")])
    return []


def extract_hof_images(doc_path: str) -> list[str]:
    """Extract HOF images from docx."""
    output_dir = WEB_DIR / "images" / "hof"
    output_dir.mkdir(parents=True, exist_ok=True)
    images = []
    
    try:
        with zipfile.ZipFile(doc_path, 'r') as z:
            for name in z.namelist():
                if name.startswith('word/media/'):
                    img_name = name.split('/')[-1]
                    data = z.read(name)
                    out_path = output_dir / img_name
                    with open(out_path, 'wb') as f:
                        f.write(data)
                    images.append(img_name)
    except Exception:
        pass
    
    return sorted(images)


def parse_transactions(doc_path: str) -> list[dict]:
    """Parse transactions document into structured seasons/weeks."""
    docx = get_docx_module()
    if not docx:
        return []
    
    doc = docx.Document(doc_path)
    seasons = []
    current_season = None
    current_week = None
    current_transaction = None
    
    LEVEL_1 = 400000
    LEVEL_2 = 800000
    LEVEL_3 = 1200000
    
    def get_indent_level(para):
        left_indent = para.paragraph_format.left_indent
        if left_indent is None:
            return 0
        if left_indent >= LEVEL_3:
            return 3
        if left_indent >= LEVEL_2:
            return 2
        if left_indent >= LEVEL_1:
            return 1
        return 0
    
    def save_transaction():
        nonlocal current_transaction
        if current_transaction and current_week:
            current_week['transactions'].append(current_transaction)
            current_transaction = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style = para.style.name if para.style else ''
        
        if style == 'Title':
            continue
        elif style == 'Heading 1':
            save_transaction()
            current_season = {'season': text, 'weeks': []}
            seasons.append(current_season)
            current_week = None
        elif style == 'Heading 2':
            save_transaction()
            if current_season:
                current_week = {'title': text, 'transactions': []}
                current_season['weeks'].append(current_week)
        else:
            if text.lower() == 'none':
                continue
            
            indent = get_indent_level(para)
            
            if indent <= 1:
                save_transaction()
                current_transaction = {'title': text, 'items': []}
            elif indent == 2:
                if current_transaction:
                    current_transaction['items'].append({'type': 'header', 'text': text})
            else:
                if current_transaction:
                    current_transaction['items'].append({'type': 'item', 'text': text})
    
    save_transaction()
    
    for season in seasons:
        season['weeks'] = [w for w in season['weeks'] if w['transactions']]
    
    return seasons


def load_transaction_log() -> list[dict]:
    """Load transactions from the JSON log file."""
    log_path = PROJECT_DIR / 'data' / 'transaction_log.json'
    if log_path.exists():
        with open(log_path) as f:
            return json.load(f).get("transactions", [])
    return []


def format_transaction_for_display(tx: dict) -> dict:
    """Format a JSON transaction into the display format."""
    items = []
    
    if tx["type"] == "trade":
        items.append({"type": "header", "text": tx.get("timestamp", "")[:10].replace("-", "/")})
        items.append({"type": "header", "text": f"To {tx.get('partner', 'Unknown')}:"})
        for player in tx.get("proposer_gives", {}).get("players", []):
            items.append({"type": "item", "text": player})
        for pick in tx.get("proposer_gives", {}).get("picks", []):
            items.append({"type": "item", "text": pick})
        items.append({"type": "header", "text": f"To {tx.get('proposer', 'Unknown')}:"})
        for player in tx.get("proposer_receives", {}).get("players", []):
            items.append({"type": "item", "text": player})
        for pick in tx.get("proposer_receives", {}).get("picks", []):
            items.append({"type": "item", "text": pick})
        
        return {
            "title": f"Trade between {tx.get('proposer', '?')} and {tx.get('partner', '?')}",
            "items": items
        }
    
    elif tx["type"] == "taxi_activation":
        items.append({"type": "header", "text": tx.get("timestamp", "")[:10].replace("-", "/")})
        items.append({"type": "header", "text": f"Activated {tx.get('activated', 'Unknown')}, released {tx.get('released', 'Unknown')}"})
        return {"title": tx.get("team", "Unknown"), "items": items}
    
    elif tx["type"] == "fa_activation":
        items.append({"type": "header", "text": tx.get("timestamp", "")[:10].replace("-", "/")})
        items.append({"type": "header", "text": f"Added {tx.get('added', 'Unknown')} from FA Pool, released {tx.get('released', 'Unknown')}"})
        return {"title": tx.get("team", "Unknown"), "items": items}
    
    return {"title": "Unknown Transaction", "items": items}


def merge_transaction_log(doc_transactions: list[dict]) -> list[dict]:
    """Merge JSON log transactions with document transactions."""
    json_transactions = load_transaction_log()
    
    if not json_transactions:
        return doc_transactions
    
    week_transactions = {}
    for tx in json_transactions:
        week = tx.get("week", 0)
        if week not in week_transactions:
            week_transactions[week] = []
        week_transactions[week].append(format_transaction_for_display(tx))
    
    current_season = None
    for season in doc_transactions:
        if "2025" in season.get("season", ""):
            current_season = season
            break
    
    if not current_season:
        current_season = {"season": "2025 Season", "weeks": []}
        doc_transactions.insert(0, current_season)
    
    for week_num, txs in week_transactions.items():
        week_title = f"Week {week_num}"
        existing_week = None
        for w in current_season["weeks"]:
            if f"Week {week_num}" in w.get("title", ""):
                existing_week = w
                break
        
        if existing_week:
            existing_week["transactions"] = txs + existing_week["transactions"]
        else:
            new_week = {"title": week_title, "transactions": txs}
            inserted = False
            for i, w in enumerate(current_season["weeks"]):
                try:
                    existing_week_num = int(''.join(filter(str.isdigit, w["title"].split()[0:2][1])) or 0)
                    if week_num > existing_week_num:
                        current_season["weeks"].insert(i, new_week)
                        inserted = True
                        break
                except (ValueError, IndexError, KeyError):
                    continue
            if not inserted:
                current_season["weeks"].append(new_week)
    
    return doc_transactions


def find_doc(name: str) -> Path:
    """Find document in docs folder or root."""
    docs_path = DOCS_DIR / name
    root_path = PROJECT_DIR / name
    return docs_path if docs_path.exists() else root_path


def export_shared():
    """Export all shared/static data."""
    ensure_dirs()
    
    print("Exporting shared data...")
    
    # Constitution
    constitution_path = find_doc("Constitution of the QPFL.docx")
    if constitution_path.exists() and get_docx_module():
        print("  - constitution.json")
        constitution = parse_constitution(str(constitution_path))
        with open(SHARED_DIR / "constitution.json", 'w') as f:
            json.dump({
                "updated_at": datetime.now(timezone.utc).isoformat(),
                "articles": constitution
            }, f, indent=2)
    
    # Hall of Fame
    hof_path = find_doc("QPFL Hall of Fame.docx")
    if hof_path.exists() and get_docx_module():
        print("  - hall_of_fame.json")
        hof = parse_hall_of_fame(str(hof_path))
        extract_hof_images(str(hof_path))
        with open(SHARED_DIR / "hall_of_fame.json", 'w') as f:
            json.dump({
                "updated_at": datetime.now(timezone.utc).isoformat(),
                **hof
            }, f, indent=2)
    
    # Banners
    print("  - banners.json")
    banners = get_banners()
    with open(SHARED_DIR / "banners.json", 'w') as f:
        json.dump({
            "updated_at": datetime.now(timezone.utc).isoformat(),
            "banners": banners
        }, f, indent=2)
    
    # Transactions
    transactions_path = find_doc("Transactions.docx")
    if transactions_path.exists() and get_docx_module():
        print("  - transactions.json")
        doc_transactions = parse_transactions(str(transactions_path))
    else:
        doc_transactions = []
    
    transactions = merge_transaction_log(doc_transactions)
    with open(SHARED_DIR / "transactions.json", 'w') as f:
        json.dump({
            "updated_at": datetime.now(timezone.utc).isoformat(),
            "seasons": transactions
        }, f, indent=2)
    
    print("Shared data exported!")


if __name__ == "__main__":
    export_shared()

